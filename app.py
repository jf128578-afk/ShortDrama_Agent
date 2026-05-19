import streamlit as st
import pandas as pd
import os
import json
import io

try:
    from openai import OpenAI
    import docx
    import openpyxl
except ImportError:
    st.error("🚨 缺少依赖库！请在下方终端输入并运行：\npip install openai python-docx openpyxl")
    st.stop()

# ==========================================
# 0. 页面基础设置 & 视觉样式
# ==========================================
st.set_page_config(page_title="爆款短剧工作台 V9.2", layout="wide", initial_sidebar_state="expanded")

st.markdown("""
    <style>
    .stTabs [data-baseweb="tab"] p { font-size: 18px !important; font-weight: 600 !important; }
    .micro-btn button { font-size: 12px !important; padding: 2px 8px !important; }
    </style>
""", unsafe_allow_html=True)

MODEL_OPTIONS = ["DeepSeek-Chat (快速/低成本版)", "DeepSeek-Reasoner (专家/深度思考版)"]

# ==========================================
# 1. 核心 AI 调用函数 (保留深度思考)
# ==========================================
def call_deepseek(prompt, model_choice, api_key):
    if not api_key:
        st.warning("⚠️ 请先在左侧边栏输入你的 DeepSeek API Key！")
        return None, None
    real_model = "deepseek-reasoner" if "Reasoner" in model_choice else "deepseek-chat"
    client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com")
    try:
        with st.spinner(f"AI 正在处理中 ({real_model})..."):
            response = client.chat.completions.create(
                model=real_model,
                messages=[{"role": "user", "content": prompt}],
                temperature=1.0 if real_model == "deepseek-chat" else None
            )
            result_text = response.choices[0].message.content
            thinking_text = getattr(response.choices[0].message, 'reasoning_content', "")
            return thinking_text, result_text
    except Exception as e:
        st.error(f"🚨 API 调用失败，请检查网络。详情：{e}")
        return None, None

# ==========================================
# 2. 导出文件生成器 (Word & Excel)
# ==========================================
def generate_word(project_name, ep_name, outline, script):
    doc = docx.Document()
    doc.add_heading(f'【{project_name}】 - {ep_name}', 0)
    
    doc.add_heading('一、 单集大纲', level=1)
    doc.add_paragraph(outline if outline else "（暂无大纲）")
    
    doc.add_heading('二、 剧本正文', level=1)
    doc.add_paragraph(script if script else "（暂无正文）")
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def generate_excel(df, project_name):
    bio = io.BytesIO()
    export_df = df[["集数", "参考_大纲骨架", "原创_大纲", "状态"]]
    export_df.to_excel(bio, index=False, engine='openpyxl')
    return bio.getvalue()

# ==========================================
# 3. 后台基建：本地项目库管理
# ==========================================
DATA_DIR = "workspace_data"
if not os.path.exists(DATA_DIR): os.makedirs(DATA_DIR)

def get_project_list():
    files = [f.replace(".json", "") for f in os.listdir(DATA_DIR) if f.endswith(".json")]
    return files if files else ["默认测试项目"]

DEFAULT_ROW = {
    "集数": "第 1 集", 
    "参考_原稿": "", "参考_清洗": "", "参考_大纲骨架": "", 
    "原创_大纲": "", "原创_正文": "", "状态": "⬜ 空白"
}

def load_project_data(project_name):
    filepath = os.path.join(DATA_DIR, f"{project_name}.json")
    if os.path.exists(filepath):
        with open(filepath, 'r', encoding='utf-8') as f: 
            df = pd.DataFrame(json.load(f))
            for key in DEFAULT_ROW.keys():
                if key not in df.columns: df[key] = ""
            return df
    return pd.DataFrame([DEFAULT_ROW])

def save_project_data(project_name, df):
    with open(os.path.join(DATA_DIR, f"{project_name}.json"), 'w', encoding='utf-8') as f:
        json.dump(df.to_dict('records'), f, ensure_ascii=False, indent=2)

def delete_project_data(project_name):
    filepath = os.path.join(DATA_DIR, f"{project_name}.json")
    if os.path.exists(filepath): os.remove(filepath)

@st.dialog("➕ 新建项目")
def create_project_dialog():
    new_proj_name = st.text_input("项目名称", label_visibility="collapsed", placeholder="例如：狂飙短剧版")
    if st.button("🚀 确认创建", type="primary", use_container_width=True):
        if new_proj_name and new_proj_name not in get_project_list():
            save_project_data(new_proj_name, pd.DataFrame([DEFAULT_ROW]))
            st.rerun()
        elif new_proj_name in get_project_list(): st.error("项目名已存在。")

@st.dialog("⚙️ 项目设置")
def manage_project_dialog(current_proj):
    st.warning(f"当前操作项目：【{current_proj}】")
    if current_proj != "默认测试项目" and st.button("🗑️彻底删除当前项目", type="primary", use_container_width=True):
        delete_project_data(current_proj)
        st.rerun()

# ==========================================
# 4. 侧边栏：配置与导航
# ==========================================
st.sidebar.title("🔑 API 配置")
# 兼容你刚才发来的版本结构
api_key_input = st.sidebar.text_input("DeepSeek API Key (sk-...)", type="password")

st.sidebar.title("📁 项目与引擎")
current_project = st.sidebar.selectbox("📂 当前打开", get_project_list())
c1, c2 = st.sidebar.columns(2)
if c1.button("➕ 新建项目", use_container_width=True): create_project_dialog()
if c2.button("⚙️ 项目设置", use_container_width=True): manage_project_dialog(current_project)

global_model = st.sidebar.selectbox("🤖 全局默认模型", MODEL_OPTIONS)
def get_model_index(): return MODEL_OPTIONS.index(global_model)

st.sidebar.divider()

st.sidebar.title("📚 剧本目录")
current_df = load_project_data(current_project)
existing_episodes = [ep for ep in current_df["集数"].dropna().tolist() if str(ep).strip() != ""]
nav_options = ["📊 全局进度与大盘"] + existing_episodes
selected_nav = st.sidebar.radio("选择工作区", nav_options)

st.sidebar.divider()

with st.sidebar.popover("💡 灵感备忘录 & AI 助理", use_container_width=True):
    st.markdown("**📝 临时备忘录** (自动保存)")
    memo = st.text_area("随时记录一闪而过的金句或设定...", value=st.session_state.get(f"memo_{current_project}", ""), height=150)
    if memo != st.session_state.get(f"memo_{current_project}", ""):
        st.session_state[f"memo_{current_project}"] = memo
    
    st.divider()
    st.markdown("**🤖 AI 智囊区**")
    chat_input = st.text_input("向 AI 提问")
    if st.button("发送提问", key="btn_chat"):
        _, res = call_deepseek(chat_input, global_model, api_key_input)
        if res: st.success(res)

# ==========================================
# 5. 主工作区：功能实现
# ==========================================
if selected_nav == "📊 全局进度与大盘":
    st.title(f"📊 【{current_project}】 全局进度与大盘")
    st.info("💡 这是您的【稳定版大盘】。在底部空白行输入集数可新增。右侧 Tab4 可将此盘直接导出为 Excel。")
    edited_df = st.data_editor(current_df, num_rows="dynamic", use_container_width=True, height=500)
    if not edited_df.equals(current_df):
        save_project_data(current_project, edited_df)
        st.rerun()

else:
    st.title(f"🎬 当前编辑：{selected_nav}")
    idx = current_df.index[current_df["集数"] == selected_nav].tolist()[0]
    ep_data = current_df.iloc[idx].to_dict()
    
    tab1, tab2, tab3, tab4 = st.tabs(["📝 1. 原稿清洗", "🦴 2. 提取参考骨架", "✍️ 3. 原创剧本撰写", "📦 4. 审阅与交付"])

    # --- TAB 1: 洗稿 ---
    with tab1:
        st.info("💡 处理参考标尺的杂质。")
        c1, c2 = st.columns(2)
        with c1: raw_text = st.text_area("贴入原始文本", value=ep_data.get("参考_原稿", ""), height=300)
        with c2: clean_val = st.text_area("清洗结果", value=ep_data.get("参考_清洗", ""), height=300)
            
        step1_model = st.selectbox("⚙️ 引擎配置", MODEL_OPTIONS, index=0, key=f"m1_{selected_nav}")
        if st.button("🚀 开始清洗"):
            prompt = f"清洗以下 OCR 剧本，去错别字理顺排版：\n\n{raw_text}"
            _, res = call_deepseek(prompt, step1_model, api_key_input)
            if res:
                current_df.at[idx, "参考_原稿"] = raw_text; current_df.at[idx, "参考_清洗"] = res
                save_project_data(current_project, current_df); st.rerun()

    # --- 🌟 完美修复的 TAB 2: 拆解骨架 🌟 ---
    with tab2:
        st.info("💡 提炼高颜值【骨架标尺】。支持手动粘贴或上传文档，勾选您需要的闪光点，AI 会带颜色高亮输出。")
        step2_model = st.selectbox("⚙️ 引擎配置", MODEL_OPTIONS, index=get_model_index(), key=f"m2_{selected_nav}")
        
        # --- 新增：闪光点提取开关 ---
        st.markdown("**✨ 附加闪光点提取（打勾后 AI 会针对性深挖）**")
        cb_col1, cb_col2, cb_col3 = st.columns(3)
        with cb_col1: need_quotes = st.checkbox("🎯 提取金句台词", value=True, key=f"cb1_{selected_nav}")
        with cb_col2: need_action = st.checkbox("💥 提取打脸/爽点动作", value=True, key=f"cb2_{selected_nav}")
        with cb_col3: need_scene = st.checkbox("🎬 提取场景调度", value=False, key=f"cb3_{selected_nav}")
        
        c1, c2 = st.columns(2)
        with c1: 
            # --- 新增：文件上传功能 ---
            uploaded_file = st.file_uploader("📁 上传剧本文档读取 (支持 TXT/Word，上传后将覆盖下方输入框)", type=["txt", "docx"], key=f"up_{selected_nav}")
            if uploaded_file is not None:
                if uploaded_file.name.endswith(".txt"):
                    file_text = uploaded_file.read().decode("utf-8")
                    ep_data["参考_清洗"] = file_text
                elif uploaded_file.name.endswith(".docx"):
                    doc = docx.Document(uploaded_file)
                    file_text = "\n".join([para.text for para in doc.paragraphs])
                    ep_data["参考_清洗"] = file_text
            
            # --- 修复：改为可编辑的文本框，再也不会“写不进去”了 ---
            input_text = st.text_area("待拆解文本（可在此直接粘贴，或使用上方上传）", value=ep_data.get("参考_清洗", ""), height=350, key=f"in2_{selected_nav}")
            if input_text != ep_data.get("参考_清洗", ""):
                current_df.at[idx, "参考_清洗"] = input_text
                save_project_data(current_project, current_df)
                
        with c2:
            extracted_html = ep_data.get("参考_大纲骨架", "")
            if extracted_html: st.markdown(f"<div style='background-color:#f8f9fa; padding:15px; border-radius:8px; border:1px solid #ddd; height:410px; overflow-y:auto;'>{extracted_html}</div>", unsafe_allow_html=True)
            else: st.write("（暂未提取，请先在左侧提供内容并点击下方拆解按钮）")
                
        if st.button("🧠 一键提取极致骨架 (带颜色高亮)", type="primary", use_container_width=True):
            if not input_text.strip():
                st.warning("⚠️ 请先在左侧填入或上传需要拆解的剧本内容！")
            else:
                # 根据复选框动态生成 Prompt
                extra_req = ""
                if need_quotes: extra_req += "\n- 单独总结出本集的【核心金句台词】"
                if need_action: extra_req += "\n- 单独提炼出本集的【核心打脸/爽点动作】"
                if need_scene: extra_req += "\n- 单独分析本集的【场景调度特点】"
                
                prompt = f"""请拆解以下短剧文本。严格按以下模块输出：【单集概述】、【事件块】、【结构分析】、【核心钩子】、【主基调】。{extra_req}
                必须使用 Markdown 和 HTML 标签进行高亮视觉排版：
                1. 核心的“强情绪台词”或“动作”，用 <span style="color:red; font-weight:bold;">文字</span> 标红。
                2. 编剧的底层意图或反转设定，用 <span style="color:green; font-weight:bold;">文字</span> 标绿。
                3. 如果有金句，用蓝色字体标出。
                请直接输出排版好的内容：\n\n{input_text}"""
                
                think, res = call_deepseek(prompt, step2_model, api_key_input)
                if res: 
                    current_df.at[idx, "参考_大纲骨架"] = res
                    current_df.at[idx, "参考_清洗"] = input_text # 同步保存左侧输入内容
                    save_project_data(current_project, current_df)
                    st.session_state[f"think_2_{selected_nav}"] = think
                    st.rerun()
                    
        if st.session_state.get(f"think_2_{selected_nav}", ""):
            with st.expander("🧠 查看 AI 拆解思考过程"): st.write(st.session_state[f"think_2_{selected_nav}"])

    # --- TAB 3: 撰写合并版 (方案 A: 单选按钮切换大纲/正文) ---
    with tab3:
        write_mode = st.radio("🔘 选择当前创作环节", ["🏗️ 阶段 A：搭原创大纲", "✍️ 阶段 B：挤牙膏写正文"], horizontal=True)
        st.divider()

        if write_mode == "🏗️ 阶段 A：搭原创大纲":
            c1, c2 = st.columns([1, 1])
            with c1:
                ref_ep = st.selectbox("👈 选择参考标尺", existing_episodes, index=existing_episodes.index(selected_nav) if selected_nav in existing_episodes else 0)
                ref_idx = current_df.index[current_df["集数"] == ref_ep].tolist()[0]
                st.markdown(f"<div style='background-color:#f0f8ff; padding:15px; border-radius:8px;'>{current_df.iloc[ref_idx]['参考_大纲骨架']}</div>", unsafe_allow_html=True)
                if st.button("📋 提取标尺至输入框"):
                    st.session_state[f"prompt_{selected_nav}"] = f"请参考以下节奏：\n{current_df.iloc[ref_idx]['参考_大纲骨架']}"
            with c2:
                step3a_model = st.selectbox("⚙️ 引擎配置", MODEL_OPTIONS, index=get_model_index(), key=f"m3a_{selected_nav}")
                user_setting = st.text_area("输入设定或标尺结构：", value=st.session_state.get(f"prompt_{selected_nav}", ""), height=100)
                outline_val = st.text_area("您的单集大纲", value=ep_data.get("原创_大纲", ""), height=250)
                
                if outline_val != ep_data.get("原创_大纲", ""):
                    current_df.at[idx, "原创_大纲"] = outline_val; save_project_data(current_project, current_df)
                    
                col_a, col_b = st.columns(2)
                if col_a.button("🤖 自动推演大纲", type="primary"):
                    prompt = f"你是一个短剧编剧。根据以下要求，写出本集【原创大纲】，不写台词：\n{user_setting}"
                    think, res = call_deepseek(prompt, step3a_model, api_key_input)
                    if res: current_df.at[idx, "原创_大纲"] = res; save_project_data(current_project, current_df); st.session_state[f"think_3a_{selected_nav}"] = think; st.rerun()
                if col_b.button("🩺 骨架诊断 (AI点评)"):
                    prompt = f"作为总监，尖锐批评以下大纲节奏和爽点不足之处，给出修改建议：\n{outline_val}"
                    _, res = call_deepseek(prompt, "DeepSeek-Reasoner", api_key_input)
                    if res: st.warning(res)
                if st.session_state.get(f"think_3a_{selected_nav}", ""):
                    with st.expander("🧠 查看推演思考过程"): st.write(st.session_state[f"think_3a_{selected_nav}"])

        else: # 写正文
            c1, c2 = st.columns([1, 1.2])
            with c1:
                st.info("您的原创大纲基准")
                st.write(ep_data.get("原创_大纲", "暂无大纲"))
                st.divider()
                step3b_model = st.selectbox("⚙️ 引擎配置 (建议专家版)", MODEL_OPTIONS, index=1, key=f"m3b_{selected_nav}")
                scene_req = st.text_area("输入分场撰写要求 (例：写男女主在饭局对峙)")
                if st.button("🔥 生成本场正文", type="primary"):
                    prompt = f"基于大纲基调，完成具体场次撰写。\n要求：{scene_req}\n格式：【场景】【动作】【台词】。必须说人话。"
                    think, res = call_deepseek(prompt, step3b_model, api_key_input)
                    if res:
                        old_text = ep_data.get("原创_正文", "")
                        current_df.at[idx, "原创_正文"] = old_text + "\n\n" + res if old_text else res
                        save_project_data(current_project, current_df)
                        st.session_state[f"think_3b_{selected_nav}"] = think; st.rerun()
                if st.session_state.get(f"think_3b_{selected_nav}", ""):
                    with st.expander("🧠 查看 AI 写这场的内心戏"): st.write(st.session_state[f"think_3b_{selected_nav}"])

            with c2:
                script_val = st.text_area("最终剧本正文区", value=ep_data.get("原创_正文", ""), height=350)
                if script_val != ep_data.get("原创_正文", ""):
                    current_df.at[idx, "原创_正文"] = script_val; save_project_data(current_project, current_df)
                
                # 🌟 强化版微操台 (带自定义指令)
                st.markdown("**🎮 台词/动作微操台**")
                mc1, mc2, mc3, mc4 = st.columns(4)
                if mc1.button("🤬 更下沉粗暴"):
                    _, res = call_deepseek(f"把台词改得更下沉粗暴：\n{script_val}", global_model, api_key_input)
                    if res: current_df.at[idx, "原创_正文"] = res; save_project_data(current_project, current_df); st.rerun()
                if mc2.button("😭 增加绿茶感"):
                    _, res = call_deepseek(f"增加绿茶语感和微表情：\n{script_val}", global_model, api_key_input)
                    if res: current_df.at[idx, "原创_正文"] = res; save_project_data(current_project, current_df); st.rerun()
                if mc3.button("💥 加肢体动作"):
                    _, res = call_deepseek(f"合理插入强烈肢体冲突描写：\n{script_val}", global_model, api_key_input)
                    if res: current_df.at[idx, "原创_正文"] = res; save_project_data(current_project, current_df); st.rerun()
                if mc4.button("🚀 融入近期热梗"):
                    _, res = call_deepseek(f"自然融入短剧热梗：\n{script_val}", global_model, api_key_input)
                    if res: current_df.at[idx, "原创_正文"] = res; save_project_data(current_project, current_df); st.rerun()
                
                # 自定义指令区
                st.markdown("**💬 自定义修改指令**")
                col_cmd, col_send = st.columns([4, 1])
                with col_cmd:
                    custom_cmd = st.text_input("具体指令", label_visibility="collapsed", placeholder="例：把这段改得隐忍一点...")
                with col_send:
                    if st.button("发送", use_container_width=True):
                        if custom_cmd:
                            _, res = call_deepseek(f"按照指令：{custom_cmd}，重新改写以下正文：\n{script_val}", global_model, api_key_input)
                            if res: current_df.at[idx, "原创_正文"] = res; save_project_data(current_project, current_df); st.rerun()

    # --- 🌟 TAB 4: 强大的导出与交付 ---
    with tab4:
        st.info("💡 生产力终点站：支持专业的 Word 与 Excel 排版导出。")
        
        col_w, col_e = st.columns(2)
        with col_w:
            st.markdown("### 📄 文本交付 (给大编剧/导演)")
            st.write("自动将您在 Tab 3 编写的【大纲】与【正文】合成标准专业排版文档。")
            word_data = generate_word(current_project, selected_nav, ep_data.get("原创_大纲", ""), ep_data.get("原创_正文", ""))
            st.download_button(
                label=f"⬇️ 导出单集 (Word .docx)", 
                data=word_data, 
                file_name=f"{current_project}_{selected_nav}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True
            )
            
        with col_e:
            st.markdown("### 📊 大盘交付 (给制片人/平台)")
            st.write("将您全局大盘中的【参考拆解骨架】与【自己写的全剧大纲】合并打包成表格。")
            excel_data = generate_excel(current_df, current_project)
            st.download_button(
                label=f"⬇️ 导出全局总进度 (Excel .xlsx)", 
                data=excel_data, 
                file_name=f"{current_project}_全局大盘归档.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                type="primary",
                use_container_width=True
            )